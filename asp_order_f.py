import tkinter as tk
import docx
from docx2pdf import convert
from tkinter import ttk
import os
import datetime
from tkinter import *
from typing import Any

dt: datetime = datetime.datetime.now()
t: datetime = dt.strftime('%d.%m.%Y')
y: datetime = dt.strftime('%Y')
trucks: list = []
names: list = []
file: str

with open("trucks.txt", "r", encoding="utf-8") as file:
    for line in file:
        trucks = list(line.split(','))
with open("drivers.txt", "r", encoding="utf-8") as file:
    for line in file:
        names = list(line.split(','))


def suggest_vehicle(event: Any) -> None:
    """
    Create Combobox object
    """
    value: Any = event.widget.get()
    if value == '':
        new_text3['values']: list = trucks
    else:
        data: list = []
        for item in trucks:
            if value.lower() in item.lower():
                data.append(item)
        new_text3['values']: list = data


def suggest_name(event: Any) -> None:
    """
    Create Combobox object
    """
    value: Any = event.widget.get()
    if value == '':
        new_text4['values']: list = names
    else:
        data: list = []
        for item in names:
            if value.lower() in item.lower():
                data.append(item)
        new_text4['values']: list = data


def add_driver() -> None:
    """
    Places new driver in existed list of drivers
    """
    new_name: Any = new_text6.get()
    with open("drivers.txt", "a", encoding="utf-8") as f:
        f.write(str(r"," + new_name))
        window.destroy()


def add_truck() -> None:
    """
    Places new truck in existed list of trucks
    """
    new_truck: Any = new_text7.get()
    with open("trucks.txt", "a", encoding="utf-8") as f:
        f.write(str(r"," + new_truck))
        window.destroy()


def change():
    """
    Changes pair time a.m. and p.m. by radiobutton, global values important because of usages in below function
    """
    global value6, value7
    if var.get() == 0:
        value6 = 'Дневная смена 08:00'
        value7 = 'День 08:00'
        
    else:
        value6 = 'Вечерняя смена 19:00'
        value7 = 'Вечер 19:00'
        


window: Tk = tk.Tk()
window.geometry('400x500')
window.title("Заказ асфальта в PDF")
window.iconbitmap("road.ico")
var: Any = IntVar()  # construct an integer variable
var.set(0)
red: Radiobutton = tk.Radiobutton(text="День 08:00", variable=var, value=0)
green: Radiobutton = tk.Radiobutton(text="Вечер 19:00", variable=var, value=1)
button: Button = tk.Button(text="Изменить", command=change)

# block clipboard
window.clipboard_clear()
window.update()
# block labels definding
label: Label = tk.Label(text='Выбрать новые значения:')
label2: Label = tk.Label(text='xx xxxxxx ' + y + ' года')
label3: Label = tk.Label(text='xx - номер заявки')
label4: Label = tk.Label(text='xxxxxxxxx - номер камаза')
label5: Label = tk.Label(text='Ф.И.О. полностью')
label6: Label = tk.Label(text='x - количество тонн')
label7: Label = tk.Label(text='Добавить нового водителя')
label8: Label = tk.Label(text='Добавить новый камаз')
# Create entryfield
new_text1: Entry = tk.Entry(window)
new_text2: Entry = tk.Entry(window)
new_text3: Any = ttk.Combobox(window, value=trucks, width=50)
new_text3.bind('<KeyRelease>', suggest_vehicle)
new_text4: Any = ttk.Combobox(window, value=names, width=50)
new_text4.bind('<KeyRelease>', suggest_name)
new_text5: Entry = tk.Entry(window)
new_text6: Entry = tk.Entry(window, width=50)
new_text7: Entry = tk.Entry(window, width=50)


def replace_text():
    """
    Create function, which replace text in exist docx file named 1
    """
    doc = docx.Document('1.docx')
    # Get new values from entry fields
    value1: object = new_text1.get()
    value2: object = new_text2.get()
    value3: object = new_text3.get()
    value4: object = new_text4.get()
    value5: object = new_text5.get()
    # Find exist text in docx
    for paragraph in doc.paragraphs:
        if '25 февраля' in paragraph.text:
            paragraph.text = paragraph.text.replace('25 февраля', value1)
        if '11' in paragraph.text:
            paragraph.text = paragraph.text.replace('11', value2)
        if 'Т079ЕО790' in paragraph.text:
            paragraph.text = paragraph.text.replace('Т079ЕО790', value3)
        if 'Иззатшоев Навбахор Алиризоевич' in paragraph.text:
            paragraph.text = paragraph.text.replace('Иззатшоев Навбахор Алиризоевич', value4)
        if '2,5 тонн' in paragraph.text:
            paragraph.text = paragraph.text.replace('2,5', value5)
        if 'Дневная смена 08:00' in paragraph.text:
            paragraph.text = paragraph.text.replace('Дневная смена 08:00', value6)

            # Find text strings in table
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '2,5 тонн' in paragraph.text:
                        paragraph.text = paragraph.text.replace('2,5', value5)
                    if 'День 08:00' in paragraph.text:
                        paragraph.text = paragraph.text.replace('День 08:00', value7)
    doc.save('2.docx')


if os.path.exists("C:/Users/ГБУ Жилищник ТС/Desktop/"):
    default_path = f"C:/Users/ГБУ Жилищник ТС/Desktop/{t}.pdf/"
else:
    default_path = "C:/Users/999.pdf/"


def convert_to_pdf():
    """
    Function which convert docx to pdf
    """
    convert("C:/425/asp_order/2.docx", default_path)
    window.destroy()


# Create action buttons
replace_button: Button = tk.Button(window, text="Заменить текст", command=replace_text)
convert_button: Button = tk.Button(window, text="Конвертировать в PDF", command=convert_to_pdf)
adding_driwer: Button = tk.Button(window, text="Добавить водителя", command=add_driver)
adding_truck: Button = tk.Button(window, text="Добавить камаз", command=add_truck)
# Place elements on the window
label.pack()
label2.pack()
new_text1.pack()
label3.pack()
new_text2.pack()
label4.pack()
new_text3.pack()
label5.pack()
new_text4.pack()
label6.pack()
new_text5.pack()
replace_button.pack()
convert_button.pack()
label7.pack()
new_text6.pack()
adding_driwer.pack()
label8.pack()
new_text7.pack()
adding_truck.pack()
red.pack()
green.pack()
button.pack()
label.pack()

window.mainloop()
